#!/usr/bin/python
# -*- coding: utf-8 -*-
# @Author: Geekwolf
# @Date:   2018-05-07 13:26:12
# @Last Modified by:   Geekwolf
# @Last Modified time: 2018-07-11 14:11:00

import ConfigParser
import cookielib
import urllib2
import urllib
import ast
import datetime,time
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Inches
from io import BytesIO
import collections
from ftplib import FTP
import os
import mimetypes
import sys
import smtplib
from email.header import Header
from email.mime.text import MIMEText
from email.mime.image import MIMEImage
from email.mime.multipart import MIMEMultipart
from email.utils import parseaddr, formataddr

config = ConfigParser.RawConfigParser()
config.read(os.path.join(os.path.dirname(os.path.abspath(__file__)),'./config.ini'))

class ZabbixGraph(object):

    def __init__(self):

        self.url = config.get('monitor', 'zbx_url')
        self.username = config.get('monitor', 'username')
        self.password = config.get('monitor', 'password')
        self.graph_url = self.url + config.get('monitor', 'graph_url')
        self.item_graph_url = self.url + config.get('monitor', 'item_graph_url')
        self.width = config.get('graph', 'width')
        self.height = config.get('graph', 'height')
        self.period = config.get('graph', 'period')
        self.temp_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)),config.get('monitor', 'temp_dir'))
	    self.log_file = os.path.join(os.path.dirname(os.path.abspath(__file__)),config.get('monitor','log_file'))
        if not os.path.exists(self.temp_dir):
            os.makedirs(self.temp_dir)
        self.urlOpener = self.GetSession()

    def GetSession(self):

        cookiejar = cookielib.CookieJar()
        urlOpener = urllib2.build_opener(urllib2.HTTPCookieProcessor(cookiejar))
        values = {"name": self.username, 'password': self.password, 'autologin': 1, "enter": 'Sign in'}
        data = urllib.urlencode(values)
        request = urllib2.Request(self.url, data)
        try:
            urlOpener.open(request, timeout=10)
            self.urlOpener = urlOpener
        except urllib2.HTTPError, e:
            print e
        return urlOpener

    def Log(self,rec):
    	t = time.strftime('%Y-%m-%d %H:%M:%S')
    	with open(self.log_file,'a') as f:
    	    f.write('{} {}'.format(str(t),str(rec)))
    	
	
    def GetRequest(self,values,url,id):
  
        _data = urllib.urlencode(values)
        request = urllib2.Request(url, _data)
        url = self.urlOpener.open(request)
        ext = mimetypes.guess_extension(url.headers['content-type'])
        imagename = '{}/{}{}'.format(self.temp_dir, str(id), ext)
        with open(imagename, 'wb') as f:
            f.write(url.read())
	return imagename

    def GetItemGraph(self,id):

	values = {'itemids': id, 'width': self.width, 'height': self.height, 'period': self.period}
        imagename = self.GetRequest(values,self.item_graph_url,id)
	return imagename

    def GetGraph(self):

        info = ast.literal_eval(config.get('graph', 'info'))
        data = collections.defaultdict(list)
        for i in info:
            values = {}
            for j in i['graphids']:
                values = {'graphid': j, 'width': self.width, 'height': self.height, 'period': self.period}
		imagename = self.GetRequest(values,self.graph_url,j)
                # image = BytesIO()
                # image.write(url.read())
                data[i['name']].append(imagename)
                # imagename = "%s/%s.png" % (self.temp_dir, str(j) + i['name'])
                # f = open(imagename, 'wb')
                # f.write(image)
        self.WriteDoc(data)

    def GetStyles(self):

        # doc = Document()
        # 在脚本打包成二进制时，需要指定default.docx路径，否则会报错
        doc = Document(docx=os.path.join(os.getcwd(), 'default.docx'))
        style_head = doc.styles.add_style('style_head', WD_STYLE_TYPE.PARAGRAPH)
        style_head.font.size = Pt(25)
        style_head.font.name = u'微软雅黑'
        style_head.font.bold = True
        style_head._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')
        style_head.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

        style_title = doc.styles.add_style('style_title', WD_STYLE_TYPE.PARAGRAPH)
        style_title.font.size = Pt(15)
        style_title.font.name = u'微软雅黑'
        style_title.font.bold = True
        style_title._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')

        sub_title = doc.styles.add_style('sub_title', WD_STYLE_TYPE.PARAGRAPH)
        sub_title.font.size = Pt(10)
        sub_title.font.name = u'微软雅黑'
        sub_title.font.bold = True
        sub_title._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')

        return doc, style_head, style_title, sub_title

    @staticmethod
    def GetYesterdayTime():

        _time = datetime.date.today() - datetime.timedelta(days=1)
        return str(_time)

    def WriteDoc(self, data):

        doc, style_head, style_title, sub_title = self.GetStyles()
        _dict = {0: '一', 1: '二', 2: '三'}
        _time = ZabbixGraph.GetYesterdayTime()
        head = doc.add_paragraph(u'zbx监控报表', style='style_head')
        sub_head = doc.add_paragraph(_time)
        sub_head.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for d in enumerate(data):
            title = '{}、{}'.format(_dict[d[0]], d[1])
            doc.add_paragraph(title.decode('utf8'), style='style_title')
            for idx, val in enumerate(data[d[1]]):
                #sub_title = u'内存' if idx%2 == 1 else u'CPU'
        		if idx >=2:
        		    sub_title = u'内存'
        		else:
        		    sub_title = u'CPU'
        		if idx%2 != 1:
                    doc.add_paragraph(sub_title, style='sub_title')
                    doc.add_picture(val, width=Inches(6.5), height=Inches(3))
        file = 'report-{}.docx'.format(('').join(_time.split('-')))
        doc.save(file)
        #如果将报表上传ftp,可以去掉注释
        #self.FtpUpload(file)

    def DelTemp(self):

        os.system('rm -rf report-* {}'.format(self.temp_dir))

    def FtpUpload(self, file):

        host = config.get('ftp', 'host')
        port = config.get('ftp', 'port')
        username = config.get('ftp', 'username')
        password = config.get('ftp', 'password')
        ftp = FTP(host=host)
        ftp.login(user=username, passwd=password)
        ftp.storbinary('STOR ' + file, open(file, 'rb'))
        ftp.quit()
        self.DelTemp()

class AlarmInfo(ZabbixGraph):

    def format(content):

        name, addr = parseaddr(content)
        return formataddr((Header(name, 'utf-8').encode(), addr))

    def Email(self):

    	smtpserver = config.get('email', 'smtpserver')
    	username = config.get('email','username')
    	password = config.get('email','password')
    	port = config.get('email','password')
    	try:
    	    smtp = smtplib.SMTP()
     	    smtp.connect(smtpserver)
    	    smtp.login(username,password)
    	    return smtp
    	except Exception as e:
    	    self.Log(str(e))

    def SendEmail(self,_info):

    	itemid = _info[2].split('|')[0]
    	imagename = self.GetItemGraph(itemid)
    	fro = config.get('email','username')
    	_content = ('<br/>').join(_info[3].split('\n'))
    	content = '{}<br><img src="cid:image1"/>'.format(_content)
    	msg = MIMEMultipart()
        #msg['From'] = '监控告警<{}>'.format(fro).decode('utf-8')
    	msg['From'] = "%s<%s>" % (Header("监控告警","utf-8"),fro)
    	msg['Subject'] = Header((': ').join(_info[2].split('|')[-2:]),'utf-8')
    	msg['To'] = _info[1]
        msg.attach(MIMEText(content,'html','utf-8'))   
        with open(imagename,'rb') as f:
       	    img = MIMEImage(f.read())
            img.add_header('Content-ID', '<image1>')
            msg.attach(img)
    	try:
    	    email = self.Email()
    	    email.sendmail(fro,_info[1],msg.as_string())
    	    email.quit()
    	except Exception as e:
    	    self.Log(str(e))

    def main(self,_info):
    	try:
    	    if len(_info) == 4:
    	        rec = '{}\t{}\n{}\n'.format(_info[1],_info[2],_info[3])
    	        self.Log(rec)
    	        self.SendEmail(_info)
    	    elif len(_info) == 2 and _info[1] == 'report':
    		    self.GetGraph()
    	except Exception as e:
    	    self.Log(str(e))	
    		

if __name__ == '__main__':

    ins = AlarmInfo()
    ins.main(sys.argv)
